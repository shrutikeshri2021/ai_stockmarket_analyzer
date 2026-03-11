"""
EDA Report Generator for AITrade
=================================
Generates Exploratory Data Analysis charts, summaries, and downloadable
PDF / Word reports from stock OHLCV DataFrames.

This module is additive — it does NOT modify any existing project logic.
"""

import io, os, datetime, tempfile
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")                       # headless backend for report generation
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# ---------------------------------------------------------------------------
# 1.  EDA CHARTS  (Plotly, shown inside Streamlit)
# ---------------------------------------------------------------------------

def generate_eda_charts(df: pd.DataFrame) -> dict:
    """Return a dict of Plotly figures keyed by chart name."""
    dates = df["Date"] if "Date" in df.columns else df.index
    charts = {}

    # 1A — Closing Price Trend
    fig_close = go.Figure()
    fig_close.add_trace(go.Scatter(
        x=dates, y=df["Close"], mode="lines", name="Close",
        line=dict(color="#00e5ff", width=2),
        fill="tozeroy", fillcolor="rgba(0,229,255,0.08)",
    ))
    fig_close.update_layout(
        title="Closing Price Trend", template="plotly_dark", height=360,
        margin=dict(l=50, r=30, t=40, b=30),
        yaxis_title="Price ($)", xaxis_title="Date",
        plot_bgcolor="#0d1117", paper_bgcolor="#0d1117",
    )
    charts["closing_price_trend"] = fig_close

    # 1B — Volume Trend
    vol_colors = ["#00e676" if df["Close"].iloc[i] >= df["Open"].iloc[i]
                  else "#ff1744" for i in range(len(df))]
    fig_vol = go.Figure()
    fig_vol.add_trace(go.Bar(x=dates, y=df["Volume"], name="Volume",
                             marker_color=vol_colors, opacity=0.75))
    fig_vol.update_layout(
        title="Volume Trend", template="plotly_dark", height=320,
        margin=dict(l=50, r=30, t=40, b=30),
        yaxis_title="Volume", xaxis_title="Date",
        plot_bgcolor="#0d1117", paper_bgcolor="#0d1117",
    )
    charts["volume_trend"] = fig_vol

    # 1C — Moving Averages (20 & 50)
    fig_ma = go.Figure()
    fig_ma.add_trace(go.Scatter(x=dates, y=df["Close"], name="Close",
                                line=dict(color="#80cbc4", width=1.5)))
    if len(df) >= 20:
        fig_ma.add_trace(go.Scatter(
            x=dates, y=df["Close"].rolling(20).mean(), name="MA 20",
            line=dict(color="#ff4081", width=2, dash="dot")))
    if len(df) >= 50:
        fig_ma.add_trace(go.Scatter(
            x=dates, y=df["Close"].rolling(50).mean(), name="MA 50",
            line=dict(color="#ffab00", width=2)))
    fig_ma.update_layout(
        title="Moving Averages (20 & 50)", template="plotly_dark", height=360,
        margin=dict(l=50, r=30, t=40, b=30),
        yaxis_title="Price ($)", xaxis_title="Date",
        plot_bgcolor="#0d1117", paper_bgcolor="#0d1117",
    )
    charts["moving_averages"] = fig_ma

    # 1D — Daily Return Distribution
    daily_returns = df["Close"].pct_change().dropna()
    fig_ret = go.Figure()
    fig_ret.add_trace(go.Histogram(
        x=daily_returns, nbinsx=60, name="Daily Returns",
        marker_color="#7c4dff", opacity=0.85,
    ))
    fig_ret.update_layout(
        title="Daily Return Distribution", template="plotly_dark", height=340,
        margin=dict(l=50, r=30, t=40, b=30),
        xaxis_title="Daily Return", yaxis_title="Frequency",
        plot_bgcolor="#0d1117", paper_bgcolor="#0d1117",
    )
    charts["daily_returns"] = fig_ret

    # 1E — Correlation Heatmap
    corr_cols = [c for c in ["Open", "High", "Low", "Close", "Volume"] if c in df.columns]
    corr = df[corr_cols].corr()
    fig_corr = go.Figure(go.Heatmap(
        z=corr.values, x=corr_cols, y=corr_cols,
        colorscale="Viridis", zmin=-1, zmax=1,
        text=np.round(corr.values, 2), texttemplate="%{text}",
        textfont=dict(color="#fff", size=13),
    ))
    fig_corr.update_layout(
        title="Correlation Heatmap", template="plotly_dark", height=420,
        margin=dict(l=60, r=30, t=40, b=60),
        plot_bgcolor="#0d1117", paper_bgcolor="#0d1117",
    )
    charts["correlation_heatmap"] = fig_corr

    return charts


# ---------------------------------------------------------------------------
# 2.  EDA SUMMARY  (text)
# ---------------------------------------------------------------------------

def generate_eda_summary(df: pd.DataFrame, ticker: str = "") -> dict:
    """Return a dict with EDA statistics and a human-readable summary string."""
    n_rows = len(df)
    dates = df["Date"] if "Date" in df.columns else df.index
    date_min = str(pd.Timestamp(dates.iloc[0]).date())
    date_max = str(pd.Timestamp(dates.iloc[-1]).date())
    year_min = pd.Timestamp(dates.iloc[0]).year
    year_max = pd.Timestamp(dates.iloc[-1]).year

    close = df["Close"]
    mean_close = close.mean()
    median_close = close.median()
    min_close = close.min()
    max_close = close.max()
    std_close = close.std()
    daily_returns = close.pct_change().dropna()
    volatility = daily_returns.std() * np.sqrt(252) * 100  # annualised %

    missing = int(df.isnull().sum().sum())
    duplicates = int(df.duplicated().sum())

    corr_cols = [c for c in ["Open", "High", "Low", "Close", "Volume"] if c in df.columns]
    corr = df[corr_cols].corr()
    vol_close_corr = corr.loc["Volume", "Close"] if "Volume" in corr.columns and "Close" in corr.index else None

    # Trend direction
    if len(close) >= 50:
        ma50_start = close.iloc[:50].mean()
        ma50_end = close.iloc[-50:].mean()
        trend = "upward" if ma50_end > ma50_start * 1.02 else ("downward" if ma50_end < ma50_start * 0.98 else "sideways")
    else:
        trend = "indeterminate (insufficient data)"

    vol_label = "low" if volatility < 20 else ("moderate" if volatility < 40 else "high")

    summary_text = (
        f"Dataset contains {n_rows:,} trading records spanning {year_min}–{year_max} "
        f"({date_min} to {date_max}).\n"
        f"The average closing price is ${mean_close:,.2f} with a median of ${median_close:,.2f}.\n"
        f"The highest price recorded is ${max_close:,.2f} and the lowest is ${min_close:,.2f}.\n"
        f"Annualised volatility is {volatility:.1f}% ({vol_label}).\n"
        f"The long-term trend direction is {trend}.\n"
    )
    if vol_close_corr is not None:
        if abs(vol_close_corr) > 0.5:
            summary_text += f"Volume and Close price show a notable correlation of {vol_close_corr:.2f}.\n"
        else:
            summary_text += f"Volume and Close price have a weak correlation of {vol_close_corr:.2f}, suggesting they move somewhat independently.\n"
    if missing == 0 and duplicates == 0:
        summary_text += "The dataset is clean — no missing values or duplicate rows detected.\n"
    else:
        summary_text += f"Data quality: {missing} missing values, {duplicates} duplicate rows.\n"

    return {
        "n_rows": n_rows,
        "date_range": f"{date_min} to {date_max}",
        "year_range": f"{year_min}–{year_max}",
        "mean_close": mean_close,
        "median_close": median_close,
        "min_close": min_close,
        "max_close": max_close,
        "std_close": std_close,
        "volatility_pct": volatility,
        "volatility_label": vol_label,
        "trend": trend,
        "missing_values": missing,
        "duplicates": duplicates,
        "vol_close_corr": vol_close_corr,
        "summary_text": summary_text,
    }


# ---------------------------------------------------------------------------
# 3.  MATPLOTLIB CHARTS  (for embedding in PDF / Word reports)
# ---------------------------------------------------------------------------

def _make_mpl_charts(df: pd.DataFrame) -> dict:
    """Create Matplotlib figures (as BytesIO PNGs) for embedding in reports."""
    dates = pd.to_datetime(df["Date"]) if "Date" in df.columns else df.index
    images = {}

    def _fig_to_bytes(fig) -> bytes:
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor="#0d1117", edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    plt_style = {"axes.facecolor": "#161b22", "figure.facecolor": "#0d1117",
                 "text.color": "#c9d1d9", "axes.labelcolor": "#c9d1d9",
                 "xtick.color": "#8b949e", "ytick.color": "#8b949e",
                 "axes.edgecolor": "#30363d", "grid.color": "#21262d"}

    with plt.rc_context(plt_style):

        # Closing Price Trend
        fig, ax = plt.subplots(figsize=(9, 3.5))
        ax.plot(dates, df["Close"], color="#00e5ff", linewidth=1.2)
        ax.fill_between(dates, df["Close"], alpha=0.08, color="#00e5ff")
        ax.set_title("Closing Price Trend", fontsize=12, fontweight="bold", color="#ffffff")
        ax.set_ylabel("Price ($)")
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
        fig.autofmt_xdate()
        ax.grid(True, alpha=0.3)
        images["closing_price_trend"] = _fig_to_bytes(fig)

        # Volume Trend
        fig, ax = plt.subplots(figsize=(9, 3))
        colors = ["#00e676" if df["Close"].iloc[i] >= df["Open"].iloc[i]
                  else "#ff1744" for i in range(len(df))]
        ax.bar(dates, df["Volume"], color=colors, width=1.0, alpha=0.75)
        ax.set_title("Volume Trend", fontsize=12, fontweight="bold", color="#ffffff")
        ax.set_ylabel("Volume")
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
        fig.autofmt_xdate()
        ax.grid(True, alpha=0.3)
        images["volume_trend"] = _fig_to_bytes(fig)

        # Moving Averages
        fig, ax = plt.subplots(figsize=(9, 3.5))
        ax.plot(dates, df["Close"], color="#80cbc4", linewidth=1, label="Close")
        if len(df) >= 20:
            ax.plot(dates, df["Close"].rolling(20).mean(), color="#ff4081",
                    linewidth=1.5, linestyle="--", label="MA 20")
        if len(df) >= 50:
            ax.plot(dates, df["Close"].rolling(50).mean(), color="#ffab00",
                    linewidth=1.5, label="MA 50")
        ax.set_title("Moving Averages (20 & 50)", fontsize=12, fontweight="bold", color="#ffffff")
        ax.set_ylabel("Price ($)")
        ax.legend(facecolor="#161b22", edgecolor="#30363d", labelcolor="#c9d1d9")
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
        fig.autofmt_xdate()
        ax.grid(True, alpha=0.3)
        images["moving_averages"] = _fig_to_bytes(fig)

        # Daily Returns Distribution
        daily_returns = df["Close"].pct_change().dropna()
        fig, ax = plt.subplots(figsize=(9, 3.5))
        ax.hist(daily_returns, bins=60, color="#7c4dff", alpha=0.85, edgecolor="#0d1117")
        ax.axvline(0, color="#ff1744", linestyle="--", linewidth=1)
        ax.set_title("Daily Return Distribution", fontsize=12, fontweight="bold", color="#ffffff")
        ax.set_xlabel("Daily Return")
        ax.set_ylabel("Frequency")
        ax.grid(True, alpha=0.3)
        images["daily_returns"] = _fig_to_bytes(fig)

        # Correlation Heatmap
        corr_cols = [c for c in ["Open", "High", "Low", "Close", "Volume"] if c in df.columns]
        corr = df[corr_cols].corr()
        fig, ax = plt.subplots(figsize=(6, 5))
        cax = ax.imshow(corr.values, cmap="viridis", vmin=-1, vmax=1, aspect="auto")
        ax.set_xticks(range(len(corr_cols)))
        ax.set_yticks(range(len(corr_cols)))
        ax.set_xticklabels(corr_cols, fontsize=10)
        ax.set_yticklabels(corr_cols, fontsize=10)
        for i in range(len(corr_cols)):
            for j in range(len(corr_cols)):
                ax.text(j, i, f"{corr.values[i, j]:.2f}", ha="center", va="center",
                        color="#ffffff", fontsize=11, fontweight="bold")
        fig.colorbar(cax, ax=ax, fraction=0.046, pad=0.04)
        ax.set_title("Correlation Heatmap", fontsize=12, fontweight="bold", color="#ffffff")
        images["correlation_heatmap"] = _fig_to_bytes(fig)

    return images


# ---------------------------------------------------------------------------
# 4.  PDF REPORT  (reportlab)
# ---------------------------------------------------------------------------

def generate_pdf_report(df: pd.DataFrame, ticker: str, company_name: str = "") -> bytes:
    """Generate a professional EDA PDF report and return the bytes."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch, mm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Image as RLImage, Table, TableStyle,
                                        PageBreak, HRFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        raise ImportError("reportlab is required for PDF generation. Install with: pip install reportlab")

    summary = generate_eda_summary(df, ticker)
    images = _make_mpl_charts(df)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=30 * mm, bottomMargin=20 * mm,
                            leftMargin=20 * mm, rightMargin=20 * mm)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"],
                                 fontSize=26, textColor=HexColor("#1a73e8"),
                                 spaceAfter=6, alignment=TA_CENTER)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"],
                                    fontSize=14, textColor=HexColor("#555555"),
                                    alignment=TA_CENTER, spaceAfter=4)
    heading_style = ParagraphStyle("HeadingCustom", parent=styles["Heading2"],
                                   fontSize=16, textColor=HexColor("#1a73e8"),
                                   spaceBefore=18, spaceAfter=8,
                                   borderWidth=1, borderColor=HexColor("#1a73e8"),
                                   borderPadding=4)
    body_style = ParagraphStyle("BodyCustom", parent=styles["Normal"],
                                fontSize=11, leading=16, spaceAfter=6,
                                textColor=HexColor("#333333"))
    insight_style = ParagraphStyle("Insight", parent=styles["Normal"],
                                   fontSize=11, leading=16, spaceAfter=4,
                                   textColor=HexColor("#2e7d32"),
                                   bulletIndent=10, leftIndent=20)

    display_name = company_name if company_name else ticker
    today_str = datetime.date.today().strftime("%B %d, %Y")

    # ---- TITLE PAGE ----
    story.append(Spacer(1, 80))
    story.append(Paragraph(f"📈 {display_name}", title_style))
    story.append(Paragraph("Exploratory Data Analysis Report", subtitle_style))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="60%", thickness=2, color=HexColor("#1a73e8"),
                             spaceAfter=12, spaceBefore=6))
    story.append(Paragraph(f"Stock Ticker: <b>{ticker}</b>", ParagraphStyle(
        "center", parent=body_style, alignment=TA_CENTER)))
    story.append(Paragraph(f"Date Generated: <b>{today_str}</b>", ParagraphStyle(
        "center2", parent=body_style, alignment=TA_CENTER)))
    story.append(Paragraph("Generated by AITrade — AI Stock Prediction System", ParagraphStyle(
        "center3", parent=body_style, alignment=TA_CENTER, textColor=HexColor("#888888"))))
    story.append(PageBreak())

    # ---- DATASET OVERVIEW ----
    story.append(Paragraph("1. Dataset Overview", heading_style))
    overview_data = [
        ["Data Source", "Yahoo Finance (via yfinance)"],
        ["Stock Ticker", ticker],
        ["Time Range", summary["date_range"]],
        ["Number of Records", f"{summary['n_rows']:,}"],
        ["Features Used", "Date, Open, High, Low, Close, Volume"],
    ]
    t = Table(overview_data, colWidths=[150, 320])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), HexColor("#e8f0fe")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("TEXTCOLOR", (0, 0), (-1, -1), HexColor("#333333")),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    # ---- DATA CLEANING ----
    story.append(Paragraph("2. Data Cleaning Summary", heading_style))
    story.append(Paragraph(f"• Missing values: <b>{summary['missing_values']}</b>", body_style))
    story.append(Paragraph(f"• Duplicate rows: <b>{summary['duplicates']}</b>", body_style))
    story.append(Paragraph("• Data preparation: Min-Max normalization applied for ML training; "
                           "raw OHLCV used for technical analysis and EDA.", body_style))
    story.append(Spacer(1, 8))

    # ---- STATISTICAL SUMMARY ----
    story.append(Paragraph("3. Statistical Summary", heading_style))
    stat_data = [
        ["Metric", "Value"],
        ["Mean Close", f"${summary['mean_close']:,.2f}"],
        ["Median Close", f"${summary['median_close']:,.2f}"],
        ["Min Close", f"${summary['min_close']:,.2f}"],
        ["Max Close", f"${summary['max_close']:,.2f}"],
        ["Std Deviation", f"${summary['std_close']:,.2f}"],
        ["Annualised Volatility", f"{summary['volatility_pct']:.1f}%"],
    ]
    t2 = Table(stat_data, colWidths=[180, 290])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1a73e8")),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 1), (0, -1), HexColor("#e8f0fe")),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t2)
    story.append(Spacer(1, 12))

    # ---- VISUAL ANALYSIS ----
    story.append(Paragraph("4. Visual Analysis", heading_style))
    chart_titles = [
        ("closing_price_trend", "Closing Price Trend"),
        ("volume_trend", "Volume Trend"),
        ("moving_averages", "Moving Averages (20 & 50)"),
        ("daily_returns", "Daily Return Distribution"),
        ("correlation_heatmap", "Correlation Heatmap"),
    ]
    for key, label in chart_titles:
        if key in images:
            story.append(Paragraph(f"<b>{label}</b>", body_style))
            img_buf = io.BytesIO(images[key])
            img = RLImage(img_buf)
            # Scale to fit page width
            max_w = 460
            aspect = img.imageWidth / img.imageHeight if img.imageHeight else 1
            img_w = min(max_w, img.imageWidth * 0.6)
            img_h = img_w / aspect
            img.drawWidth = img_w
            img.drawHeight = img_h
            story.append(img)
            story.append(Spacer(1, 10))

    # ---- KEY INSIGHTS ----
    story.append(PageBreak())
    story.append(Paragraph("5. Key Insights", heading_style))

    insights = [
        f"The long-term trend direction is <b>{summary['trend']}</b>.",
        f"Annualised volatility is <b>{summary['volatility_pct']:.1f}%</b> "
        f"(<b>{summary['volatility_label']}</b>), indicating the degree of price fluctuation.",
    ]
    if summary["vol_close_corr"] is not None:
        vc = summary["vol_close_corr"]
        if abs(vc) > 0.5:
            insights.append(f"Volume and Close price show a <b>notable correlation ({vc:.2f})</b>, "
                            "suggesting price moves are often accompanied by volume spikes.")
        else:
            insights.append(f"Volume and Close price have a <b>weak correlation ({vc:.2f})</b>, "
                            "meaning volume alone is not a strong predictor of price direction.")
    if summary["max_close"] and summary["min_close"]:
        price_range_pct = ((summary["max_close"] - summary["min_close"]) / summary["min_close"]) * 100
        insights.append(f"The stock traded in a range of <b>${summary['min_close']:,.2f}</b> to "
                        f"<b>${summary['max_close']:,.2f}</b> — a <b>{price_range_pct:.1f}%</b> spread.")
    _mv = summary['missing_values']
    _dp = summary['duplicates']
    if _mv == 0 and _dp == 0:
        insights.append("The dataset is <b>clean</b> (no missing values or duplicates).")
    else:
        insights.append(f"The dataset has <b>{_mv} missing values</b> and <b>{_dp} duplicates</b>.")

    for ins in insights:
        story.append(Paragraph(f"• {ins}", insight_style))

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cccccc"),
                             spaceAfter=8, spaceBefore=8))
    story.append(Paragraph(
        f"<i>Report generated by AITrade v2.0 on {today_str}. "
        "Educational purposes only — not financial advice.</i>",
        ParagraphStyle("footer", parent=body_style, fontSize=9,
                       textColor=HexColor("#999999"), alignment=TA_CENTER),
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# 5.  WORD (.docx) REPORT  (python-docx)
# ---------------------------------------------------------------------------

def generate_word_report(df: pd.DataFrame, ticker: str, company_name: str = "") -> bytes:
    """Generate a professional EDA Word (.docx) report and return the bytes."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx is required for Word generation. Install with: pip install python-docx")

    summary = generate_eda_summary(df, ticker)
    images = _make_mpl_charts(df)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    display_name = company_name if company_name else ticker
    today_str = datetime.date.today().strftime("%B %d, %Y")

    # ---- TITLE PAGE ----
    for _ in range(4):
        doc.add_paragraph("")
    title = doc.add_heading(f"{display_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Exploratory Data Analysis Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    info = doc.add_paragraph(f"Stock Ticker: {ticker}\nDate Generated: {today_str}\n"
                             "Generated by AITrade — AI Stock Prediction System")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ---- DATASET OVERVIEW ----
    doc.add_heading("1. Dataset Overview", level=1)
    overview_rows = [
        ("Data Source", "Yahoo Finance (via yfinance)"),
        ("Stock Ticker", ticker),
        ("Time Range", summary["date_range"]),
        ("Number of Records", f"{summary['n_rows']:,}"),
        ("Features Used", "Date, Open, High, Low, Close, Volume"),
    ]
    tbl = doc.add_table(rows=len(overview_rows), cols=2, style="Light Grid Accent 1")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(overview_rows):
        tbl.cell(i, 0).text = k
        tbl.cell(i, 1).text = v
    doc.add_paragraph("")

    # ---- DATA CLEANING ----
    doc.add_heading("2. Data Cleaning Summary", level=1)
    doc.add_paragraph(f"Missing values: {summary['missing_values']}")
    doc.add_paragraph(f"Duplicate rows: {summary['duplicates']}")
    doc.add_paragraph("Data preparation: Min-Max normalization applied for ML training; "
                      "raw OHLCV used for technical analysis and EDA.")

    # ---- STATISTICAL SUMMARY ----
    doc.add_heading("3. Statistical Summary", level=1)
    stat_rows = [
        ("Metric", "Value"),
        ("Mean Close", f"${summary['mean_close']:,.2f}"),
        ("Median Close", f"${summary['median_close']:,.2f}"),
        ("Min Close", f"${summary['min_close']:,.2f}"),
        ("Max Close", f"${summary['max_close']:,.2f}"),
        ("Std Deviation", f"${summary['std_close']:,.2f}"),
        ("Annualised Volatility", f"{summary['volatility_pct']:.1f}%"),
    ]
    tbl2 = doc.add_table(rows=len(stat_rows), cols=2, style="Light Grid Accent 1")
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(stat_rows):
        tbl2.cell(i, 0).text = k
        tbl2.cell(i, 1).text = v
        if i == 0:
            for cell in tbl2.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
    doc.add_paragraph("")

    # ---- VISUAL ANALYSIS ----
    doc.add_heading("4. Visual Analysis", level=1)
    chart_titles = [
        ("closing_price_trend", "Closing Price Trend"),
        ("volume_trend", "Volume Trend"),
        ("moving_averages", "Moving Averages (20 & 50)"),
        ("daily_returns", "Daily Return Distribution"),
        ("correlation_heatmap", "Correlation Heatmap"),
    ]
    for key, label in chart_titles:
        if key in images:
            doc.add_paragraph(label, style="Heading 3")
            img_stream = io.BytesIO(images[key])
            doc.add_picture(img_stream, width=Inches(5.8))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")

    # ---- KEY INSIGHTS ----
    doc.add_page_break()
    doc.add_heading("5. Key Insights", level=1)
    insights = [
        f"The long-term trend direction is {summary['trend']}.",
        f"Annualised volatility is {summary['volatility_pct']:.1f}% "
        f"({summary['volatility_label']}), indicating the degree of price fluctuation.",
    ]
    if summary["vol_close_corr"] is not None:
        vc = summary["vol_close_corr"]
        if abs(vc) > 0.5:
            insights.append(f"Volume and Close price show a notable correlation ({vc:.2f}), "
                            "suggesting price moves are often accompanied by volume spikes.")
        else:
            insights.append(f"Volume and Close price have a weak correlation ({vc:.2f}), "
                            "meaning volume alone is not a strong predictor of price direction.")
    if summary["max_close"] and summary["min_close"]:
        price_range_pct = ((summary["max_close"] - summary["min_close"]) / summary["min_close"]) * 100
        insights.append(f"The stock traded in a range of ${summary['min_close']:,.2f} to "
                        f"${summary['max_close']:,.2f} — a {price_range_pct:.1f}% spread.")
    insights.append(
        "The dataset is clean (no missing values or duplicates)."
        if summary["missing_values"] == 0 and summary["duplicates"] == 0
        else f"The dataset has {summary['missing_values']} missing values and {summary['duplicates']} duplicates."
    )
    for ins in insights:
        doc.add_paragraph(f"• {ins}")

    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"Report generated by AITrade v2.0 on {today_str}. "
        "Educational purposes only — not financial advice."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
